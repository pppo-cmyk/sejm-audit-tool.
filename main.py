python
import requests
import io
import re
import logging
import zipfile
import os
import sys
import time
import subprocess
import numpy as np
import concurrent.futures
import pandas as pd
from datetime import datetime
from pypdf import PdfReader
from pypdf.errors import FileNotDecryptedError
from pdf2image import convert_from_bytes
from unidecode import unidecode
from thefuzz import fuzz
from paddleocr import PaddleOCR
from docx import Document
import openpyxl
import xlrd

# ==============================================================================
# ‚öôÔ∏è KONFIGURACJA TOTALNA (PROJECT TOTAL RECALL)
# ==============================================================================

# System dependencies check
def check_system_dependencies():
    """Check for required system dependencies."""
    print("üîç [SYSTEM CHECK] Sprawdzanie zale≈ºno≈õci systemowych...")
    
    missing = []
    
    # Check for poppler-utils (provides pdftoppm, pdftocairo)
    try:
        subprocess.run(['pdftoppm', '-v'], stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=False)
    except FileNotFoundError:
        missing.append('poppler-utils')
    
    # Check for libgl1 (OpenGL library)
    if sys.platform.startswith('linux'):
        try:
            result = subprocess.run(['ldconfig', '-p'], stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, check=False)
            if 'libGL.so' not in result.stdout:
                missing.append('libgl1')
        except Exception:
            # If ldconfig fails, warn anyway
            missing.append('libgl1 (nie mo≈ºna sprawdziƒá)')
    
    if missing:
        print(f"‚ö†Ô∏è  [SYSTEM CHECK] BRAKUJƒÑCE ZALE≈ªNO≈öCI: {', '.join(missing)}")
        print(f"‚ö†Ô∏è  [SYSTEM CHECK] Zainstaluj u≈ºywajƒÖc: sudo apt-get install {' '.join(missing)}")
        print(f"‚ö†Ô∏è  [SYSTEM CHECK] Program mo≈ºe nie dzia≈Çaƒá poprawnie bez tych pakiet√≥w!")
        return False
    else:
        print("‚úÖ [SYSTEM CHECK] Wszystkie wymagane zale≈ºno≈õci systemowe sƒÖ zainstalowane.")
        return True

check_system_dependencies()

TERMS = [9, 10]
API_URL = "https://api.sejm.gov.pl/sejm"
OUTPUT_DIR = "sejm_audit_output"
SAVE_INTERVAL_SECONDS = 300  # Zapis co 5 minut

# High-resolution DPI for OCR
PDF_DPI = 300  # High quality scan

# WEBSHARE PROXY CONFIGURATION
# Set these environment variables: WEBSHARE_PROXY_HOST, WEBSHARE_PROXY_PORT, WEBSHARE_PROXY_USER, WEBSHARE_PROXY_PASS
PROXY_HOST = os.getenv('WEBSHARE_PROXY_HOST', '')
PROXY_PORT = os.getenv('WEBSHARE_PROXY_PORT', '')
PROXY_USER = os.getenv('WEBSHARE_PROXY_USER', '')
PROXY_PASS = os.getenv('WEBSHARE_PROXY_PASS', '')

# Proxy configuration for requests
PROXIES = None
if PROXY_HOST and PROXY_PORT:
    if PROXY_USER and PROXY_PASS:
        proxy_url = f"http://{PROXY_USER}:{PROXY_PASS}@{PROXY_HOST}:{PROXY_PORT}"
    else:
        proxy_url = f"http://{PROXY_HOST}:{PROXY_PORT}"
    PROXIES = {
        'http': proxy_url,
        'https': proxy_url
    }
    # Log proxy info without exposing credentials
    if PROXY_USER:
        masked_user = f"{PROXY_USER[:2]}***" if len(PROXY_USER) > 2 else "***"
        print(f"üåê [PROXY] U≈ºywam Webshare proxy: {masked_user}@{PROXY_HOST}:{PROXY_PORT}")
    else:
        print(f"üåê [PROXY] U≈ºywam Webshare proxy: {PROXY_HOST}:{PROXY_PORT}")
else:
    print("‚ö†Ô∏è [PROXY] Brak konfiguracji proxy - u≈ºywam bezpo≈õredniego po≈ÇƒÖczenia")

# S≈ÅOWNIK RYZYKA - MILITARY & DEFENSE FOCUS
SEMANTIC_TRIGGERS = {
    "MILITARY_DEFENSE": [
        "wojsko", "czolg", "amunicja", "f-35", "f35", "uzbrojenie",
        "obrona narodowa", "sily zbrojne", "zolnierz", "weteran",
        "modernizacja armii", "kontrakt zbrojeniowy", "himars",
        "rakieta", "zakup broni", "sprzet wojskowy", "system obrony",
        "my≈õliwiec", "mysliwiec", "czolgi", "pancerz", "artyleria",
        "wojska specjalne", "nsz", "dow", "dowodztwo", "batalion",
        "brygada", "dywizja", "regiment", "kompania", "pluton",
        "amw", "abw", "skw", "sww", "wywiad wojskowy", "kontrwywiad wojskowy",
        "rosomak", "krab", "borsuk", "narew", "wisla", "homar",
        "patriot", "piorun", "grom", "thunder", "javelin", "bayraktar",
        "m1 abrams", "k2", "fa-50", "apache", "black hawk"
    ]
}

if not os.path.exists(OUTPUT_DIR): os.makedirs(OUTPUT_DIR)

# ==============================================================================
# üöÄ OCR INITIALIZATION (CPU MODE - PREVENTS SEGFAULTS)
# ==============================================================================
# Security Note: Using PaddlePaddle 3.0.0+ to avoid CVEs in versions <= 2.6.0
# The application does not use vulnerable functions (paddle.vision.ops.read_file, 
# paddle.utils.download._wget_download) and only uses the safe PaddleOCR API.
print("‚ö° [OCR INIT] Start silnika PaddleOCR (CPU Mode)...")
try:
    GLOBAL_OCR_ENGINE = PaddleOCR(
        use_angle_cls=True,
        lang='pl',
        use_gpu=False,          # CPU mode to prevent segmentation faults
        enable_mkldnn=True,     # Enable Intel MKL-DNN acceleration for CPU
        show_log=False
    )
    print("‚úÖ [OCR INIT] Gotowy. Tryb: HEAVY AUDIT MODE - Full OCR with CPU acceleration.")
except Exception as e:
    print(f"‚ùå [OCR INIT] B≈ÇƒÖd inicjalizacji PaddleOCR: {e}")
    print("‚ö†Ô∏è  Sprawd≈∫ czy wszystkie zale≈ºno≈õci sƒÖ zainstalowane.")
    raise RuntimeError(f"B≈ÇƒÖd OCR: {e}")

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(message)s')

# ==============================================================================
# üõ†Ô∏è NARZƒòDZIA POMOCNICZE
# ==============================================================================

def get_roman(n):
    val = [1000, 900, 500, 400, 100, 90, 50, 40, 10, 9, 5, 4, 1]
    syb = ["M", "CM", "D", "CD", "C", "XC", "L", "XL", "X", "IX", "V", "IV", "I"]
    roman_num = ''
    i = 0
    while n > 0:
        for _ in range(n // val[i]):
            roman_num += syb[i]
            n -= val[i]
        i += 1
    return roman_num

def index_to_char(n):
    return chr(65 + n) if n < 26 else f"Z{n}"

def save_batch_to_disk(rows, batch_idx):
    if not rows: return
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"{OUTPUT_DIR}/audit_part_{batch_idx}_{timestamp}.csv"
    df = pd.DataFrame(rows)
    cols = ["TREE_ID", "STATUS_SKANU", "DRZEWO STRUKTURY", "Nazwa Pliku", "Link", 
            "RYZYKO", "Alerty", "Autor", "Data Pliku", "S≈Çowa"]
    for c in cols:
        if c not in df.columns: df[c] = ""
    df[cols].to_csv(filename, index=False, sep=';', encoding='utf-8-sig')
    print(f"üíæ [AUTO-SAVE] Zapisano partiƒô {batch_idx}: {filename} ({len(rows)} rekord√≥w)")

def extract_metadata(content, ext):
    """Wydobywa metadane z pliku (autor, data)."""
    metadata = {"Autor": "?", "Data": "?"}
    try:
        if ext == 'pdf':
            reader = PdfReader(io.BytesIO(content))
            if reader.metadata:
                metadata["Autor"] = reader.metadata.get('/Author', '?') or '?'
                creation_date = reader.metadata.get('/CreationDate', '?')
                if creation_date and creation_date != '?':
                    # Parse PDF date format (D:YYYYMMDDHHmmss)
                    if isinstance(creation_date, str) and creation_date.startswith('D:'):
                        try:
                            date_str = creation_date[2:10]  # YYYYMMDD
                            metadata["Data"] = f"{date_str[:4]}-{date_str[4:6]}-{date_str[6:8]}"
                        except Exception:
                            metadata["Data"] = creation_date
                    else:
                        metadata["Data"] = str(creation_date)
        elif ext in ['docx', 'doc']:
            doc = Document(io.BytesIO(content))
            if doc.core_properties:
                metadata["Autor"] = doc.core_properties.author or '?'
                if doc.core_properties.created:
                    metadata["Data"] = doc.core_properties.created.strftime('%Y-%m-%d')
    except Exception:
        pass
    return metadata

def robust_request(url, retries=3, timeout=120):
    """Pobieranie z obs≈ÇugƒÖ Rate Limit (429) i proxy failures - exponential backoff with retries."""
    base_delay = 2
    for attempt in range(retries):
        try:
            resp = requests.get(url, timeout=timeout, proxies=PROXIES)
            
            if resp.status_code == 429:  # Rate Limit
                wait = base_delay * (2 ** attempt)  # Exponential backoff
                print(f"üõë Rate Limit (429) - pr√≥ba {attempt + 1}/{retries}. Czekam {wait}s...")
                time.sleep(wait)
                continue
            
            if resp.status_code == 200:
                return resp
            
            # Other errors - retry with backoff
            if attempt < retries - 1:
                wait = base_delay * (2 ** attempt)
                print(f"‚ö†Ô∏è  HTTP {resp.status_code} - pr√≥ba {attempt + 1}/{retries}. Czekam {wait}s...")
                time.sleep(wait)
                continue
            
            return resp
            
        except requests.exceptions.ProxyError as e:
            if attempt < retries - 1:
                wait = base_delay * (2 ** attempt)
                print(f"üîå Proxy Error - pr√≥ba {attempt + 1}/{retries}. Czekam {wait}s...")
                time.sleep(wait)
            else:
                print(f"‚ùå Proxy failed after {retries} attempts: {e}")
                return None
                
        except requests.exceptions.Timeout as e:
            if attempt < retries - 1:
                wait = base_delay * (2 ** attempt)
                print(f"‚è±Ô∏è  Timeout - pr√≥ba {attempt + 1}/{retries}. Czekam {wait}s...")
                time.sleep(wait)
            else:
                print(f"‚ùå Timeout after {retries} attempts: {e}")
                return None
                
        except Exception as e:
            if attempt < retries - 1:
                wait = base_delay * (2 ** attempt)
                print(f"‚ö†Ô∏è  Error: {e} - pr√≥ba {attempt + 1}/{retries}. Czekam {wait}s...")
                time.sleep(wait)
            else:
                print(f"‚ùå Failed after {retries} attempts: {e}")
                return None
    
    return None

# ==============================================================================
# üß† FORENSIC SCANNER
# ==============================================================================

class ForensicScanner:
    def __init__(self, file_bytes, filename):
        self.file_bytes = file_bytes
        self.filename = filename
        self.ext = filename.split('.')[-1].lower()
        self.risk = 0
        self.vectors = []
        self.alerts = []
        
        self.visual_text = ""  # OCR z obrazka (GPU)
        self.logic_text = ""   # Tekst z kodu pliku

    def ocr_cpu(self, images):
        """Process images using CPU-based OCR"""
        text = ""
        for img in images:
            try:
                img_array = np.array(img)
                res = GLOBAL_OCR_ENGINE.ocr(img_array, cls=True)
                if res and res[0]:
                    text += " ".join([line[1][0] for line in res[0]]) + " "
            except Exception:
                pass
        return text

    def scan_pdf(self):
        try:
            # FULL OCR MODE - Scan every page visually using pdf2image and PaddleOCR
            # Do NOT use simple text extraction
            
            # Check if encrypted and get reader instance
            reader = None
            try:
                reader = PdfReader(io.BytesIO(self.file_bytes))
                if reader.is_encrypted:
                    try:
                        reader.decrypt('')
                    except Exception:
                        self.alerts.append("üîí ZABLOKOWANE HAS≈ÅEM")
                        self.risk += 10
                        return
            except FileNotDecryptedError:
                self.alerts.append("üîí ZABLOKOWANE HAS≈ÅEM")
                self.risk += 10
                return
            except Exception:
                pass  # Continue with OCR even if PDF reading fails

            # VISUAL LAYER - Render and OCR ALL pages at high DPI (300)
            print(f"  üî¨ [OCR] Skanowanie wizualne: {self.filename}")
            images = convert_from_bytes(
                self.file_bytes,
                dpi=PDF_DPI,          # High resolution 300 DPI
                fmt='jpeg',
                thread_count=8,
                use_pdftocairo=True
            )
            self.visual_text = self.ocr_cpu(images)
            
            # For forensic analysis, we still want logic text for comparison
            # but visual text is primary
            # Reuse reader instance if available
            if reader is not None:
                try:
                    for page in reader.pages:
                        self.logic_text += (page.extract_text() or "") + " "
                except Exception:
                    pass  # If text extraction fails, we still have OCR

        except FileNotDecryptedError:
            self.alerts.append("üîí ZABLOKOWANE HAS≈ÅEM")
            self.risk += 10
        except Exception as e:
            self.alerts.append(f"PDF Error: {str(e)}")

    def scan_docx(self):
        try:
            doc = Document(io.BytesIO(self.file_bytes))
            self.logic_text += " ".join([p.text for p in doc.paragraphs])
            for t in doc.tables:
                for r in t.rows:
                    for c in r.cells: self.logic_text += c.text + " "
            
            # Obrazki w Wordzie
            with zipfile.ZipFile(io.BytesIO(self.file_bytes)) as z:
                media = [f for f in z.namelist() if f.startswith('word/media/')]
                if media:
                    from PIL import Image
                    pil_imgs = []
                    for m in media:
                        with z.open(m) as f:
                            try:
                                pil_imgs.append(Image.open(f).convert('RGB'))
                            except Exception:
                                pass
                    if pil_imgs:
                        self.visual_text += self.ocr_cpu(pil_imgs)
                        self.alerts.append("[SKAN W WORDZIE]")
        except Exception:
            pass

    def scan_excel(self):
        try:
            # Excel traktujemy jako logiczny
            df_dict = pd.read_excel(io.BytesIO(self.file_bytes), sheet_name=None)
            for sheet_name, df in df_dict.items():
                self.logic_text += f" [Arkusz: {sheet_name}] {df.to_string()} "
        except Exception as e:
            self.alerts.append(f"Excel Error: {e}")

    def analyze_results(self):
        clean_visual = unidecode(self.visual_text).lower()
        clean_logic = unidecode(self.logic_text).lower()
        
        # ≈ÅƒÖczymy do szukania trigger√≥w
        combined_text = clean_visual + " " + clean_logic
        clean_combined = re.sub(r'[^a-z0-9\s]', '', combined_text)

        found_cats = set()
        
        # SZUKANIE S≈Å√ìW - MILITARY & DEFENSE ONLY
        for cat, terms in SEMANTIC_TRIGGERS.items():
            for term in terms:
                term_clean = unidecode(term).lower()
                # Fuzzy match
                if term_clean in clean_combined or fuzz.partial_ratio(term_clean, clean_combined) > 90:
                    self.vectors.append(term)
                    found_cats.add(cat)
                    self.risk += 3  # Higher risk score for military content

        # FORENSIC DIFF (POR√ìWNANIE WARSTW - TYLKO DLA PDF)
        if self.ext == 'pdf':
            for vec in self.vectors:
                in_logic = vec in clean_logic
                in_visual = vec in clean_visual
                
                # A. INJECTION (Bia≈Çy tekst)
                if in_logic and not in_visual:
                    self.alerts.append(f"‚ö†Ô∏è INJECTION (Tylko w kodzie): '{vec}'")
                    self.risk += 5
                
                # B. DEEP RIDER (Tylko na obrazie)
                if in_visual and not in_logic:
                    self.alerts.append(f"üëÅÔ∏è DEEP RIDER (Tylko na obrazie): '{vec}'")
                    self.risk += 5

        # Bonus for finding military content
        if "MILITARY_DEFENSE" in found_cats:
            self.alerts.append("üéØ MILITARY & DEFENSE CONTENT DETECTED")

        return min(self.risk, 10)

    def run(self):
        if self.ext == 'pdf': self.scan_pdf()
        elif self.ext in ['docx', 'doc']: self.scan_docx()
        elif self.ext in ['xlsx', 'xls']: self.scan_excel()
        else:
            try:
                self.logic_text = self.file_bytes.decode('utf-8', errors='ignore')
            except Exception:
                pass
            
        return self.analyze_results()

# ==============================================================================
# üå≥ WORKER (REKURENCJA ZIP)
# ==============================================================================

def process_file_content(content, filename, file_id, visual_tree, url):
    rows = []
    ext = filename.split('.')[-1].lower()
    
    # OBS≈ÅUGA ARCHIW√ìW (ZIP)
    if ext == 'zip':
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as z:
                rows.append({
                    "TREE_ID": file_id, "STATUS_SKANU": "OK (ZIP)",
                    "DRZEWO STRUKTURY": f"{visual_tree} üì¶ {filename}",
                    "Nazwa Pliku": filename, "Link": url, "RYZYKO": 0, "Alerty": "Rozpakowano w locie", "S≈Çowa": ""
                })
                
                for i, zip_file_name in enumerate(z.namelist()):
                    if zip_file_name.endswith('/'): continue
                    sub_content = z.read(zip_file_name)
                    sub_id = f"{file_id}.{i+1}"
                    sub_tree = visual_tree.replace("‚îî‚îÄ‚îÄ", "    ‚îî‚îÄ‚îÄ")
                    
                    rows.extend(process_file_content(
                        sub_content, zip_file_name, sub_id, f"{sub_tree} ‚Ü™Ô∏è", "wewn_zip"
                    ))
            return rows
        except Exception:
            pass 

    # PLIK POJEDYNCZY
    row = {
        "TREE_ID": file_id, "STATUS_SKANU": "OK",
        "DRZEWO STRUKTURY": f"{visual_tree} üìÑ {filename}",
        "Nazwa Pliku": filename, "Link": url,
        "RYZYKO": 0, "Alerty": "", "Autor": "?", "Data Pliku": "?", "S≈Çowa": ""
    }

    try:
        m = extract_metadata(content, ext)
        row["Autor"] = m["Autor"]
        row["Data Pliku"] = m["Data"]
        
        scanner = ForensicScanner(content, filename)
        risk = scanner.run()
        
        row["RYZYKO"] = risk
        row["S≈Çowa"] = ", ".join(list(set(scanner.vectors)))
        row["Alerty"] = " | ".join(scanner.alerts)
        
    except Exception as e:
        row["STATUS_SKANU"] = f"SCAN ERROR: {str(e)}"

    return [row]

def worker_process(proc, term, proc_idx):
    rows = []
    process_status = "OK"
    roman_id = get_roman(proc_idx)
    
    # NAG≈Å√ìWEK PROCESU
    rows.append({
        "TREE_ID": f"{roman_id}", "STATUS_SKANU": "...",
        "DRZEWO STRUKTURY": f"üìÇ [{proc.get('num', '?')}] {proc['title'][:150]}...",
        "Nazwa Pliku": "", "Link": f"https://sejm.gov.pl/Sejm{term}.nsf/przebieg.xsp?id={proc['num']}",
        "RYZYKO": "", "Alerty": "", "Autor": "", "Data Pliku": "", "S≈Çowa": ""
    })

    prints = proc.get('prints', [])
    for p_i, print_nr in enumerate(prints, 1):
        print_id = f"{roman_id}.{p_i}"
        rows.append({
            "TREE_ID": print_id, "STATUS_SKANU": "",
            "DRZEWO STRUKTURY": f"    ‚îú‚îÄ‚îÄ üìÅ Druk {print_nr}",
            "Nazwa Pliku": "", "Link": "", "RYZYKO": "", "Alerty": "", "Autor": "", "Data Pliku": "", "S≈Çowa": ""
        })

        try:
            meta_resp = robust_request(f"{API_URL}/term{term}/prints/{print_nr}")
            if not meta_resp or meta_resp.status_code != 200:
                rows[-1]["STATUS_SKANU"] = "API ERROR"
                continue
                
            attachments = meta_resp.json().get('attachments', [])
            for f_i, att in enumerate(attachments):
                file_id = f"{print_id}.{index_to_char(f_i)}"
                url = f"{API_URL}/term{term}/prints/{print_nr}/{att}"
                visual_tree = "        ‚îî‚îÄ‚îÄ"
                
                file_resp = robust_request(url)
                if file_resp and file_resp.status_code == 200:
                    file_rows = process_file_content(file_resp.content, att, file_id, visual_tree, url)
                    rows.extend(file_rows)
                else:
                    rows.append({
                        "TREE_ID": file_id, "STATUS_SKANU": "DOWNLOAD ERROR",
                        "DRZEWO STRUKTURY": f"{visual_tree} ‚ùå {att}", "Nazwa Pliku": att, "Link": url
                    })

        except Exception as e:
            process_status = f"ERROR: {str(e)}"

    rows[0]["STATUS_SKANU"] = process_status
    return rows

def get_all_processes(term):
    try:
        return requests.get(f"{API_URL}/term{term}/processes", timeout=60, proxies=PROXIES).json()
    except Exception:
        return []

# ==============================================================================
# üìù SAMPLE REPORT GENERATOR
# ==============================================================================

def generate_sample_report():
    """Generate a sample report file on startup to show expected output format."""
    sample_file = "RAPORT_PROBNY.txt"
    timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    
    sample_content = f"""
‚ïî‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïó
‚ïë                    SEJM AUDIT TOOL - RAPORT PR√ìBNY                         ‚ïë
‚ïë                    Heavy Audit Mode - Military & Defense                   ‚ïë
‚ïö‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïù

Data wygenerowania: {timestamp}
Tryb skanowania: FULL OCR (pdf2image + PaddleOCR)
Rozdzielczo≈õƒá: 300 DPI
Kategoria: MILITARY & DEFENSE

‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ

PRZYK≈ÅADOWE ZNALEZISKA (DUMMY DATA):

[1] Proces IX.1234 - Ustawa o modernizacji Si≈Ç Zbrojnych
    ‚îú‚îÄ‚îÄ Druk 1234
    ‚îÇ   ‚îî‚îÄ‚îÄ üìÑ ustawa_modernizacja.pdf
    ‚îÇ       ‚îú‚îÄ‚îÄ Ryzyko: 8/10
    ‚îÇ       ‚îú‚îÄ‚îÄ Wykryte s≈Çowa: wojsko, czo≈Çg, amunicja, F-35, uzbrojenie
    ‚îÇ       ‚îú‚îÄ‚îÄ Alerty: üéØ MILITARY & DEFENSE CONTENT DETECTED
    ‚îÇ       ‚îú‚îÄ‚îÄ Autor: Ministerstwo Obrony Narodowej
    ‚îÇ       ‚îî‚îÄ‚îÄ Data: 2024-06-15

[2] Proces X.5678 - Ustawa o zakupie sprzƒôtu obronnego
    ‚îú‚îÄ‚îÄ Druk 5678
    ‚îÇ   ‚îî‚îÄ‚îÄ üìÑ kontrakt_f35.pdf
    ‚îÇ       ‚îú‚îÄ‚îÄ Ryzyko: 9/10
    ‚îÇ       ‚îú‚îÄ‚îÄ Wykryte s≈Çowa: F-35, himars, rakieta, modernizacja armii
    ‚îÇ       ‚îú‚îÄ‚îÄ Alerty: üéØ MILITARY & DEFENSE CONTENT DETECTED | üëÅÔ∏è DEEP RIDER
    ‚îÇ       ‚îú‚îÄ‚îÄ Autor: MON
    ‚îÇ       ‚îî‚îÄ‚îÄ Data: 2024-08-22

[3] Proces X.9012 - Ustawa o wsparciu weteran√≥w
    ‚îú‚îÄ‚îÄ Druk 9012
    ‚îÇ   ‚îî‚îÄ‚îÄ üìÑ weterani_2024.pdf
    ‚îÇ       ‚îú‚îÄ‚îÄ Ryzyko: 6/10
    ‚îÇ       ‚îú‚îÄ‚îÄ Wykryte s≈Çowa: ≈ºo≈Çnierz, weteran, wojsko
    ‚îÇ       ‚îú‚îÄ‚îÄ Alerty: üéØ MILITARY & DEFENSE CONTENT DETECTED
    ‚îÇ       ‚îú‚îÄ‚îÄ Autor: Biuro Legislacyjne
    ‚îÇ       ‚îî‚îÄ‚îÄ Data: 2024-09-10

‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ

KONFIGURACJA SYSTEMU:
‚úì PaddleOCR: CPU Mode (use_gpu=False, enable_mkldnn=True)
‚úì Webshare Proxy: {'Aktywny' if PROXIES else 'Nieaktywny'}
‚úì ThreadPoolExecutor: 8 wƒÖtk√≥w (parallel processing)
‚úì Retry mechanism: 3 pr√≥by z exponential backoff
‚úì PDF DPI: 300 (high resolution)

‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ

UWAGI:
‚Ä¢ To jest tylko przyk≈Çadowy raport pokazujƒÖcy format wyj≈õciowy
‚Ä¢ Rzeczywiste wyniki bƒôdƒÖ zapisane w plikach CSV w folderze '{OUTPUT_DIR}'
‚Ä¢ Ka≈ºdy plik PDF jest skanowany wizualnie u≈ºywajƒÖc OCR
‚Ä¢ Wykrywane sƒÖ s≈Çowa kluczowe z kategorii MILITARY & DEFENSE
‚Ä¢ System automatycznie zapisuje wyniki co 5 minut

‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ

S≈ÅOWA KLUCZOWE (MILITARY & DEFENSE):
wojsko, czo≈Çg, amunicja, F-35, uzbrojenie, obrona narodowa, si≈Çy zbrojne,
≈ºo≈Çnierz, weteran, modernizacja armii, kontrakt zbrojeniowy, HIMARS, rakieta,
zakup broni, sprzƒôt wojskowy, system obrony, my≈õliwiec, pancerz, artyleria,
wojska specjalne, Rosomak, Krab, Borsuk, Narew, Wis≈Ça, Homar, Patriot,
Piorun, Grom, Thunder, Javelin, Bayraktar, M1 Abrams, K2, FA-50, Apache,
Black Hawk

‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ

Aby rozpoczƒÖƒá prawdziwe skanowanie, uruchom: python main.py

‚ïö‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïù
"""
    
    try:
        with open(sample_file, 'w', encoding='utf-8') as f:
            f.write(sample_content)
        print(f"üìã [SAMPLE REPORT] Wygenerowano przyk≈Çadowy raport: {sample_file}")
        print(f"üìã [SAMPLE REPORT] Otw√≥rz plik aby zobaczyƒá oczekiwany format wyj≈õciowy.")
    except Exception as e:
        print(f"‚ö†Ô∏è  [SAMPLE REPORT] Nie uda≈Ço siƒô wygenerowaƒá raportu: {e}")

# ==============================================================================
# üèÅ MAIN LOOP
# ==============================================================================

def main():
    print("=== SEJM HEAVY AUDIT MODE (MILITARY & DEFENSE SCANNER) ===")
    print("=== Full OCR with High Resolution (300 DPI) ===")
    
    # Generate sample report on startup
    generate_sample_report()
    print()
    
    tasks = []
    global_idx = 1
    for term in TERMS:
        procs = get_all_processes(term)
        print(f"Kadencja {term}: Znaleziono {len(procs)} proces√≥w.")
        for p in procs:
            tasks.append((p, term, global_idx))
            global_idx += 1
            
    print(f"Start pracy. Wyniki co 5 minut w folderze '{OUTPUT_DIR}'.")
    print(f"U≈ºywam {8} wƒÖtk√≥w do r√≥wnoleg≈Çego przetwarzania PDF√≥w.")
    
    buffer_rows = []
    last_save_time = time.time()
    batch_counter = 1
    
    # Use 8 threads for parallel processing (5-10 range as specified)
    with concurrent.futures.ThreadPoolExecutor(max_workers=8) as executor:
        future_to_proc = {executor.submit(worker_process, t[0], t[1], t[2]): t[0]['num'] for t in tasks}
        
        completed = 0
        total = len(tasks)
        
        for future in concurrent.futures.as_completed(future_to_proc):
            completed += 1
            proc_num = future_to_proc[future]
            try:
                res = future.result()
                if res:
                    buffer_rows.extend(res)
                    if completed % 10 == 0:
                        print(f"[{completed}/{total}] Przetworzono proces {proc_num}")
            except Exception as e:
                print(f"B≈ÇƒÖd procesu {proc_num}: {e}")

            if time.time() - last_save_time >= SAVE_INTERVAL_SECONDS:
                save_batch_to_disk(buffer_rows, batch_counter)
                buffer_rows = []
                batch_counter += 1
                last_save_time = time.time()

    if buffer_rows: save_batch_to_disk(buffer_rows, "FINAL")
    print("‚úÖ KONIEC PRACY.")

if __name__ == "__main__":
    main()
